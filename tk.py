#coding:utf-8
from Tkinter import * 
from PIL import Image, ImageTk
import tkFileDialog
import sys
import random
import primes
import copy
import docx2txt
import docx
class BlumBlumShub(object):
    def getPrime(self, bits):
        while True:
            p = primes.bigppr(bits)
            if p & 3 == 3:
                return p

    def generateN(self, bits):
        """
        This generates the "n value" for use in the Blum-Blum-Shub algorithm.
       
        bits - The number of bits of security
        """
    
        p = self.getPrime(bits/2)
        while 1:
            q = self.getPrime(bits/2)
            # make sure p != q (almost always true, but just in case, check)
            if p != q:
                return p * q    

    def __init__(self, bits):
        """
        Constructor, specifing bits for n.
         
        bits - number of bits
        """        
        self.n = self.generateN(bits)
        # print "n set to " + repr(self.n)
        length = self.bitLen(self.n)
        seed = random.getrandbits(length)
        self.setSeed(seed)  
    def setSeed(self, seed):
        """
        Sets or resets the seed value and internal state.
         
        seed -The new seed
        """
    
        self.state = seed % self.n
    
    def bitLen(self, x):
        " Get the bit lengh of a positive number" 
        assert x > 0
        q = 0 
        while x: 
            q += 1 
            x >>= 1 
        return q     

    def next(self, numBits):
        "Returns up to numBit random bits"
        
        result = 0
        for i in xrange(numBits):
            self.state = (self.state**2) % self.n
            result = (result << 1) | (self.state&1)
        
        return result    
if __name__ == "__main__":

    bbs = BlumBlumShub(128);
    def pgcd(a,b):
        if b==0:
            return a
        else:
            r=a%b
        return pgcd(b,r)
    def modinverse2(a,b):
        a = a % b
        x = 1
        while x < b:
            if (a*x) % b == 1:
                return x
            x = x+1
    def test(view,x):
        view['text']=x
    def decrypt():
        doc = docx.Document(field['text'])
        msg = doc.paragraphs[0].text
        s = str(field2['text'][13:field['text'].find(')')]).split(',')
        msg = msg.split(',')
        plain = [chr((int(char) ** int(s[1])) % int(s[0])) for char in msg]
        doc2 = docx.Document()
        paragra = doc2.add_paragraph('')
        paragra.add_run(''.join(plain))
        doc2.save(field['text'])
    def encrypt():
        doc = docx.Document(field['text'])
        msg = doc.paragraphs[0].text
        s = str(field1['text'][14:field['text'].find(')')]).split(',')
        cipher = [(ord(char) ** int(s[1])) % int(s[0]) for char in msg]
        cipher = str(cipher).replace('[','')
        cipher = str(cipher).replace('L','')
        cipher = str(cipher).replace(']','')
        cipher = str(cipher).replace(' ','')
        print cipher
        doc2 = docx.Document()
        paragra = doc2.add_paragraph('')
        paragra.add_run(str(cipher))
        doc2.save(field['text'])
    def fic(img):
        file = tkFileDialog.askopenfile(parent=fenetre,mode='rb',title='Choose a file')
        field['text']=file.name
        field['image']=img
    def isPrime2(n):
        if n==2 or n==3: return True
        if n%2==0 or n<2: return False
        for i in range(3,int(n**0.5)+1,2):   # only odd numbers
            if n%i==0:
                return False    

        return True
    def generate():
        for i in xrange (5000000):
            w = bbs.next(8)
            if isPrime2(w):
		        print(w);break
        for i in xrange (5000000):
            z = bbs.next(8)
            if isPrime2(z):
		        print(z);break
        n = w*z
        x = (w-1)*(z-1)
        for i in xrange (5000000):
            e = bbs.next(8)
            if pgcd(x,e)==1:
		        print(e);break
        print "X = ",x
        print e
        m = ord("m")
        print m
        c = pow(m,e,n)
        d = modinverse2(e,x)

        m = pow(c,d,n)
        print chr(m)
        test(field1,"Clé Public : ({0},{1})".format(n,e))
        test(field2,"Clé Privé : ({0},{1})".format(n,d))
    fenetre = Tk()
    fenetre.geometry("500x300")
    fenetre.title('RSA')
    fenetre.resizable(0, 0)
    img = ImageTk.PhotoImage(file='decrypt.png')
    img2 = ImageTk.PhotoImage(file='file.png')
    img3 = ImageTk.PhotoImage(file='keys.png')
    img4 = ImageTk.PhotoImage(file='exit.png')
    img5 = ImageTk.PhotoImage(file='word.png')
    bouton=Button(fenetre,image=img3, width=130,height=30,text="Génération des Clés",command=generate,compound=LEFT)
    bouton.place(x=50,y=20)
    #bouton.pack()
    field1 = Label(fenetre,text="")
    field1.place(x=100,y=60)
    field2 = Label(fenetre,text="")
    field2.place(x=100,y=80)
    field = Label(fenetre,text="",compound=LEFT)
    field.place(x=50,y=150)
    bouton4 = Button(fenetre, image=img2, width=130,height=30,text="Choisir un fichier",command=lambda: fic(img5),compound=LEFT)
    bouton4.place(x=50,y=110)
    bouton5 = Button(fenetre, image=img, width=130,height=30,text="Crypter",command=encrypt,compound=LEFT)
    bouton5.place(x=50,y=190)
    bouton6 = Button(fenetre, image=img, width=130,height=30,text="Dérypter",command=decrypt,compound=LEFT)
    bouton6.place(x=230,y=190)
    bouton=Button(fenetre, image=img4, width=130,height=30,text="Quitter", command=fenetre.quit,compound=LEFT)
    bouton.place(x=20,y=250)
#label.pack()

#label['text']="testkiw"
    
    
    #print "Generating 10 numbers"
    
	#p = w
	#q = z
    #n = p*q
    #print n
    fenetre.mainloop()
